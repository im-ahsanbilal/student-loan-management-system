import csv
import json
import mimetypes
import os
import secrets
from datetime import date
from datetime import datetime
from functools import wraps
from io import BytesIO, StringIO
from pathlib import Path

from flask import abort, current_app
from flask_login import current_user, login_required
from flask_mail import Message
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Image as ReportImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import func
from werkzeug.utils import secure_filename

from directory_service import managed_department_names
from extensions import db, mail
from models import Department, Document, LoanApplication, Notification, VerificationLog


PENDING_STATUSES = ["Pending", "HOD Verification", "Finance Verification", "Admin Decision", "Correction Required", "Resubmitted"]
WORKFLOW_STEPS = [
    "Application Submitted",
    "HOD Review",
    "Finance Review",
    "Admin Decision",
    "Loan Tracking",
]


def allowed_file(filename):
    if "." not in filename:
        return False
    extension = filename.rsplit(".", 1)[1].lower()
    return extension in current_app.config["ALLOWED_EXTENSIONS"]


def role_required(*roles):
    def decorator(view_func):
        @wraps(view_func)
        @login_required
        def wrapped_view(*args, **kwargs):
            if current_user.role not in roles:
                abort(403)
            return view_func(*args, **kwargs)

        return wrapped_view

    return decorator


def generate_unique_loan_id():
    while True:
        loan_id = f"SL-{secrets.token_hex(4).upper()}"
        existing = LoanApplication.query.filter_by(loan_id=loan_id).first()
        if not existing:
            return loan_id


def application_requires_correction(application):
    return application.status == "Correction Required" or (
        application.status == "Pending" and application.current_stage == "Pending Correction"
    )


def send_otp_email(user, otp, purpose):
    if purpose == "registration":
        subject = "Student Loan System OTP Verification"
        headline = "Verify your student account"
    else:
        subject = "Student Loan System Password Reset OTP"
        headline = "Reset your password"

    text_body = (
        f"{headline}\n\n"
        f"Dear {user.full_name},\n"
        f"Your OTP is {otp}. It expires in {current_app.config['OTP_EXPIRY_MINUTES']} minutes.\n"
        "If you did not request this code, please ignore this email."
    )
    html_body = (
        f"<p>Dear {user.full_name},</p>"
        f"<p>Your OTP is <strong>{otp}</strong>. It expires in "
        f"{current_app.config['OTP_EXPIRY_MINUTES']} minutes.</p>"
        "<p>If you did not request this code, please ignore this email.</p>"
    )

    if current_app.config.get("MAIL_SUPPRESS_SEND") or not current_app.config.get("MAIL_USERNAME"):
        current_app.logger.warning("OTP for %s (%s): %s", user.email, purpose, otp)
        return False

    try:
        message = Message(subject=subject, recipients=[user.email])
        message.body = text_body
        message.html = html_body
        mail.send(message)
        return True
    except Exception as exc:
        current_app.logger.error("Email send failed for %s: %s", user.email, exc)
        current_app.logger.warning("Fallback OTP for %s (%s): %s", user.email, purpose, otp)
        return False


def create_notification(user_id, title, message, application_id=None):
    notification = Notification(
        user_id=user_id,
        title=title,
        message=message,
        application_id=application_id,
    )
    db.session.add(notification)
    return notification


def log_verification(application_id, actor_id, stage, action, remarks):
    log = VerificationLog(
        application_id=application_id,
        actor_id=actor_id,
        stage=stage,
        action=action,
        remarks=remarks,
    )
    db.session.add(log)
    return log


def resolve_managed_upload_path(file_path):
    if not file_path:
        return None
    try:
        upload_root = Path(current_app.config["UPLOAD_FOLDER"]).resolve()
        resolved_path = Path(file_path).resolve()
    except OSError:
        return None
    if upload_root != resolved_path and upload_root not in resolved_path.parents:
        return None
    if not resolved_path.exists():
        return None
    return str(resolved_path)


def build_correction_request_message(application, selected_labels, remarks):
    message_lines = [f"Please update the following items for application {application.loan_id}:"]
    for label in selected_labels:
        message_lines.append(f"- {label}")
    message_lines.append(f"Notes: {remarks}")
    return "\n".join(message_lines)


def save_application_document(file_storage, application, document_type):
    filename = secure_filename(file_storage.filename)
    if not allowed_file(filename):
        raise ValueError("Unsupported file format.")

    _, extension = os.path.splitext(filename)
    folder_name = os.path.join(current_app.config["UPLOAD_FOLDER"], application.loan_id)
    os.makedirs(folder_name, exist_ok=True)

    safe_type = document_type.lower().replace(" ", "_")
    stored_filename = f"{safe_type}_{secrets.token_hex(8)}{extension.lower()}"
    file_path = os.path.join(folder_name, stored_filename)
    file_storage.save(file_path)

    existing_document = Document.query.filter_by(
        application_id=application.id,
        document_type=document_type,
    ).first()

    if existing_document:
        existing_path = resolve_managed_upload_path(existing_document.file_path)
        if existing_path:
            try:
                os.remove(existing_path)
            except OSError:
                current_app.logger.warning("Could not delete old file %s", existing_document.file_path)
        document = existing_document
        document.original_filename = filename
        document.stored_filename = stored_filename
        document.file_path = file_path
        document.mime_type = mimetypes.guess_type(filename)[0]
        document.uploaded_at = datetime.utcnow()
    else:
        document = Document(
            application_id=application.id,
            document_type=document_type,
            original_filename=filename,
            stored_filename=stored_filename,
            file_path=file_path,
            mime_type=mimetypes.guess_type(filename)[0],
        )
        db.session.add(document)

    document.review_status = "Accepted"
    document.review_remarks = None
    document.reviewed_at = None
    document.reviewed_by_id = None
    return document


def save_profile_photo(file_storage, user):
    if not file_storage or not file_storage.filename:
        raise ValueError("Passport-size photo is required.")

    filename = secure_filename(file_storage.filename)
    _, extension = os.path.splitext(filename)
    extension = extension.lower()
    if extension not in {".jpg", ".jpeg", ".png"}:
        raise ValueError("Passport-size photo must be a JPG or PNG file.")

    stream = file_storage.stream
    current_position = stream.tell()
    stream.seek(0, os.SEEK_END)
    file_size = stream.tell()
    stream.seek(current_position)
    if file_size > current_app.config["MAX_PROFILE_PHOTO_SIZE"]:
        raise ValueError("Passport-size photo must be 2 MB or smaller.")

    folder_name = os.path.join(current_app.config["UPLOAD_FOLDER"], "profile_photos")
    os.makedirs(folder_name, exist_ok=True)
    stored_filename = f"user_{user.id}_{secrets.token_hex(8)}{extension}"
    file_path = os.path.join(folder_name, stored_filename)

    existing_photo_path = resolve_managed_upload_path(user.profile_photo_path)
    if existing_photo_path:
        try:
            os.remove(existing_photo_path)
        except OSError:
            current_app.logger.warning("Could not delete old profile photo %s", user.profile_photo_path)

    file_storage.stream.seek(0)
    file_storage.save(file_path)
    user.profile_photo_path = file_path
    user.profile_photo_original_name = filename


def user_can_access_application(application):
    if current_user.role == "student":
        return application.student_id == current_user.id
    if current_user.role == "hod":
        return application.department in managed_department_names(current_user)
    if current_user.role in {"finance", "admin"}:
        return True
    return False


def enforce_application_access(application):
    if not user_can_access_application(application):
        abort(403)


def application_progress_index(application):
    if application_requires_correction(application):
        return 3
    if application.status in {"Pending", "HOD Verification"}:
        return 1
    if application.status == "Resubmitted":
        return 2
    if application.status == "Finance Verification":
        return 2
    if application.status == "Admin Decision":
        return 3
    if application.status in {"Approved", "Rejected"}:
        return 4
    return 0


def workflow_snapshot(application):
    steps = [
        {
            "label": "Application Submitted",
            "state": "complete",
            "detail": f"Submitted on {application.submitted_at.strftime('%d %b %Y')}",
        },
        {"label": "HOD Review", "state": "pending", "detail": "Awaiting department verification"},
        {"label": "Finance Review", "state": "pending", "detail": "Awaiting finance verification"},
        {"label": "Admin Decision", "state": "pending", "detail": "Awaiting final decision"},
        {"label": "Loan Tracking", "state": "pending", "detail": "Loan tracking will appear after approval"},
    ]

    if application.hod_verified():
        steps[1] = {
            "label": "HOD Review",
            "state": "complete",
            "detail": application.hod_comments or _hod_validation_summary(application),
        }
    elif application.status == "HOD Verification":
        steps[1] = {"label": "HOD Review", "state": "current", "detail": "Currently with HOD"}
    elif application.status == "Rejected" and application.rejected_by_stage == "HOD Verification":
        steps[1] = {
            "label": "HOD Review",
            "state": "rejected",
            "detail": application.last_rejection_reason() or "Not verified by HOD",
        }

    if application.finance_verified():
        steps[2] = {
            "label": "Finance Review",
            "state": "complete",
            "detail": application.finance_comments or _finance_validation_summary(application),
        }
    elif application.status == "Finance Verification" or (
        application.status == "Resubmitted" and application.current_stage == "Finance Verification"
    ):
        steps[2] = {"label": "Finance Review", "state": "current", "detail": "Currently with finance"}
    elif application.status == "Rejected" and application.rejected_by_stage == "Finance Verification":
        steps[2] = {
            "label": "Finance Review",
            "state": "rejected",
            "detail": application.last_rejection_reason() or "Not verified by finance",
        }

    if application.admin_status == "Approved" or application.status == "Approved":
        steps[3] = {
            "label": "Admin Decision",
            "state": "complete",
            "detail": application.admin_comments or "Loan approved by admin",
        }
    elif application_requires_correction(application):
        steps[3] = {
            "label": "Admin Decision",
            "state": "current",
            "detail": "Admin requested corrections before final approval.",
        }
    elif application.status == "Admin Decision":
        steps[3] = {"label": "Admin Decision", "state": "current", "detail": "Awaiting admin action"}
    elif application.status == "Rejected" and application.rejected_by_stage == "Admin Decision":
        steps[3] = {
            "label": "Admin Decision",
            "state": "rejected",
            "detail": application.last_rejection_reason() or "Rejected by admin",
        }

    if application.status == "Approved":
        repayment_text = (
            application.repayment_deadline.strftime("%d %b %Y")
            if application.repayment_deadline
            else "To be announced"
        )
        disbursement_text = (
            application.expected_disbursement_date.strftime("%d %b %Y")
            if application.expected_disbursement_date
            else "Pending schedule"
        )
        steps[4] = {
            "label": "Loan Tracking",
            "state": "complete",
            "detail": f"Disbursement: {application.disbursement_status} on {disbursement_text} | Repayment: {repayment_text}",
        }
    elif application.status == "Rejected":
        steps[4] = {
            "label": "Loan Tracking",
            "state": "pending",
            "detail": "Loan tracking is unavailable because the application was rejected",
        }

    return steps


def build_tracking_timeline(application):
    timeline = [
        {
            "title": "Application submitted",
            "timestamp": application.submitted_at,
            "state": "complete",
            "detail": f"Student submitted {application.loan_id} with required documents.",
        }
    ]

    if application.hod_verified_at:
        timeline.append(
            {
                "title": f"HOD {application.hod_status.lower()}",
                "timestamp": application.hod_verified_at,
                "state": "complete" if application.hod_verified() else "rejected",
                "detail": application.hod_comments or _hod_validation_summary(application),
            }
        )
    else:
        timeline.append(
            {
                "title": "Awaiting HOD review",
                "timestamp": None,
                "state": "current" if application.status == "HOD Verification" else "pending",
                "detail": "The assigned HOD will review departmental eligibility.",
            }
        )

    if application.finance_verified_at:
        timeline.append(
            {
                "title": f"Finance {application.finance_status.lower()}",
                "timestamp": application.finance_verified_at,
                "state": "complete" if application.finance_verified() else "rejected",
                "detail": application.finance_comments or _finance_validation_summary(application),
            }
        )
    else:
        timeline.append(
            {
                "title": "Awaiting finance review",
                "timestamp": None,
                "state": "current" if application.status == "Finance Verification" or (application.status == "Resubmitted" and application.current_stage == "Finance Verification") else "pending",
                "detail": "Finance will verify scholarship overlap and eligibility.",
            }
        )

    if application.admin_decided_at:
        timeline.append(
            {
                "title": f"Admin {application.admin_status.lower()}",
                "timestamp": application.admin_decided_at,
                "state": "complete" if application.admin_status == "Approved" else "rejected",
                "detail": application.admin_comments or "Final admin decision recorded.",
            }
        )
    else:
        timeline.append(
            {
                "title": "Awaiting admin decision",
                "timestamp": None,
                "state": "current" if application.status == "Admin Decision" else "pending",
                "detail": "Admin can finalize only after both HOD and finance verifications are complete.",
            }
        )

    if application_requires_correction(application):
        timeline.append(
            {
                "title": "Correction required",
                "timestamp": application.updated_at,
                "state": "current",
                "detail": application.correction_summary() or "Admin requested corrections to the application details or documents.",
            }
        )

    if application.status == "Resubmitted":
        timeline.append(
            {
                "title": "Submitted after correction",
                "timestamp": application.last_resubmitted_at,
                "state": "complete",
                "detail": "Student submitted the corrected application back into the review workflow.",
            }
        )

    if application.status == "Approved":
        timeline.append(
            {
                "title": "Eligible for disbursement",
                "timestamp": application.admin_decided_at,
                "state": "complete",
                "detail": "Finance has been notified that the student can proceed to loan disbursement.",
            }
        )
        if application.expected_disbursement_date:
            timeline.append(
                {
                    "title": "Disbursement scheduled",
                    "timestamp": application.expected_disbursement_date,
                    "state": "complete",
                    "detail": f"Loan will be disbursed to {application.disbursement_iban or application.student_iban}.",
                }
            )
        if application.disbursement_sent_at:
            timeline.append(
                {
                    "title": "Disbursement sent",
                    "timestamp": application.disbursement_sent_at,
                    "state": "complete",
                    "detail": "Finance marked the loan amount as sent to the student.",
                }
            )
        if application.disbursement_received_at:
            timeline.append(
                {
                    "title": "Disbursement received",
                    "timestamp": application.disbursement_received_at,
                    "state": "complete",
                    "detail": "Student receipt of the loan amount was confirmed.",
                }
            )
        if application.repayment_deadline:
            timeline.append(
                {
                    "title": "Repayment deadline",
                    "timestamp": application.repayment_deadline,
                    "state": "pending",
                    "detail": "Student repayment is due by this deadline.",
                }
            )

    if application.status == "Rejected":
        timeline.append(
            {
                "title": "Application rejected",
                "timestamp": application.updated_at,
                "state": "rejected",
                "detail": application.last_rejection_reason() or "A reviewer rejected the application.",
            }
        )

    timeline.sort(key=lambda item: (item["timestamp"] is None, _normalized_timeline_timestamp(item["timestamp"]) or datetime.max))
    return timeline


def clear_rejection_state(application):
    application.rejected_by_stage = None
    application.rejection_reason = None


def mark_application_rejected(application, stage, comments):
    application.status = "Rejected"
    application.current_stage = "Rejected"
    application.rejected_by_stage = stage
    application.rejection_reason = comments
    application.disbursement_status = "Pending"


def reset_application_for_resubmission(application, *, status="HOD Verification", current_stage="HOD Verification"):
    application.status = status
    application.current_stage = current_stage
    application.hod_status = "Pending"
    application.finance_status = "Pending"
    application.admin_status = "Pending"
    application.scholarship_status = "Not Checked"
    application.hod_cgpa_verified = False
    application.hod_department_verified = False
    application.finance_fee_defaulter = None
    application.finance_scholarship_availing = None
    application.hod_comments = None
    application.finance_comments = None
    application.admin_comments = None
    application.hod_verified_at = None
    application.finance_verified_at = None
    application.admin_decided_at = None
    application.loan_amount_issued = None
    application.disbursement_iban = None
    application.expected_disbursement_date = None
    application.repayment_deadline = None
    application.disbursement_status = "Pending"
    application.disbursement_sent_at = None
    application.disbursement_received_at = None
    application.resubmission_count = (application.resubmission_count or 0) + 1
    application.last_resubmitted_at = datetime.utcnow()
    application.field_corrections_json = None
    clear_rejection_state(application)


def application_ready_for_admin_decision(application):
    return application.approvals_complete()


def build_report_payload(status_filter=None, department_filter=None):
    query = LoanApplication.query
    if status_filter:
        query = query.filter(LoanApplication.status == status_filter)
    if department_filter:
        query = query.filter(LoanApplication.department == department_filter)

    applications = query.order_by(LoanApplication.submitted_at.desc()).all()
    pending_count = query.filter(LoanApplication.status.in_(PENDING_STATUSES)).count()
    summary = {
        "total_applications": query.count(),
        "approved_loans": query.filter(LoanApplication.status == "Approved").count(),
        "rejected_loans": query.filter(LoanApplication.status == "Rejected").count(),
        "pending_applications": pending_count,
    }

    department_query = db.session.query(
        LoanApplication.department,
        func.count(LoanApplication.id).label("total"),
    )
    if status_filter:
        department_query = department_query.filter(LoanApplication.status == status_filter)
    if department_filter:
        department_query = department_query.filter(LoanApplication.department == department_filter)
    department_rows = department_query.group_by(LoanApplication.department).order_by(LoanApplication.department.asc()).all()

    status_rows = [
        ("Pending", query.filter(LoanApplication.status == "Pending").count()),
        ("HOD Verification", query.filter(LoanApplication.status == "HOD Verification").count()),
        ("Finance Verification", query.filter(LoanApplication.status == "Finance Verification").count()),
        ("Admin Decision", query.filter(LoanApplication.status == "Admin Decision").count()),
        ("Approved", query.filter(LoanApplication.status == "Approved").count()),
        ("Rejected", query.filter(LoanApplication.status == "Rejected").count()),
    ]

    coverage_rows = (
        Department.query.order_by(Department.category.asc(), Department.name.asc()).all()
        if not department_filter
        else Department.query.filter_by(name=department_filter).all()
    )

    return {
        "applications": applications,
        "summary": summary,
        "department_rows": department_rows,
        "status_rows": status_rows,
        "coverage_rows": coverage_rows,
        "filters": {
            "status": status_filter or "",
            "department": department_filter or "",
        },
        "generated_at": datetime.utcnow(),
    }


def generate_application_pdf(application):
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=24, leftMargin=24, topMargin=28, bottomMargin=24)
    styles = _report_styles()
    elements = []

    elements.append(Paragraph("Student Loan Application Record", styles["Title"]))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(f"Loan ID: {application.loan_id}", styles["Heading2"]))
    elements.append(Paragraph(f"Generated on {datetime.utcnow().strftime('%d %b %Y %I:%M %p')}", styles["Meta"]))
    elements.append(Spacer(1, 12))

    student_photo_path = None
    if application.student and application.student.profile_photo_path:
        student_photo_path = resolve_managed_upload_path(application.student.profile_photo_path)
    if student_photo_path:
        try:
            photo = ReportImage(student_photo_path)
            max_width = 90
            max_height = 110
            width_scale = max_width / photo.drawWidth
            height_scale = max_height / photo.drawHeight
            scale = min(width_scale, height_scale, 1)
            photo.drawWidth = photo.drawWidth * scale
            photo.drawHeight = photo.drawHeight * scale
            elements.append(Paragraph("Student Passport Photo", styles["Heading3"]))
            elements.append(Spacer(1, 4))
            elements.append(photo)
            elements.append(Spacer(1, 10))
        except Exception:
            elements.append(Paragraph("Student passport photo preview unavailable.", styles["BodyText"]))
            elements.append(Spacer(1, 10))

    student_rows = [
        ["Field", "Value"],
        ["Student Name", application.student_name],
        ["Roll Number", application.roll_number],
        ["Father Name", application.father_name],
        ["Phone Number", application.phone_number or "-"],
        ["CNIC", application.cnic],
        ["CGPA", f"{float(application.cgpa):.2f}" if application.cgpa is not None else "-"],
        ["Department", application.department],
        ["Program", application.program_level or "BS"],
        ["Semester", application.semester],
    ]
    elements.append(_styled_table(student_rows, widths=[150, 360]))
    elements.append(Spacer(1, 10))

    loan_rows = [
        ["Loan Detail", "Value"],
        ["Loan Amount Requested", _format_currency(application.loan_amount_requested)],
        ["Family Income", _format_currency(application.family_income)],
        ["Bank Name", application.bank_name or "-"],
        ["IBAN", application.iban or application.student_iban or "-"],
        ["Account Holder Relation", application.account_holder_relation or application.relation_to_student or "Self"],
        ["Issued Amount", _format_currency(application.loan_amount_issued) if application.loan_amount_issued else "Pending"],
        ["Disbursement IBAN", application.disbursement_iban or "Pending"],
        ["Disbursement Tracking", application.disbursement_status or "Pending"],
        [
            "Disbursement Date",
            application.expected_disbursement_date.strftime("%d %b %Y") if application.expected_disbursement_date else "Pending",
        ],
        [
            "Disbursement Sent At",
            application.disbursement_sent_at.strftime("%d %b %Y %I:%M %p") if application.disbursement_sent_at else "Pending",
        ],
        [
            "Disbursement Received At",
            application.disbursement_received_at.strftime("%d %b %Y %I:%M %p") if application.disbursement_received_at else "Pending",
        ],
        [
            "Repayment Deadline",
            application.repayment_deadline.strftime("%d %b %Y") if application.repayment_deadline else "Pending",
        ],
        ["Reason for Loan", application.reason_for_loan],
    ]
    elements.append(_styled_table(loan_rows, widths=[150, 360]))
    elements.append(Spacer(1, 10))

    status_rows = [
        ["Verification Stage", "Status / Comments"],
        ["Application Status", application.status],
        ["Current Stage", application.current_stage],
        ["HOD Verification", f"{application.hod_status} | {application.hod_comments or 'N/A'}"],
        ["HOD CGPA Check", _yes_no_unknown(application.hod_cgpa_verified, false_label="Not Confirmed")],
        ["HOD Department Check", _yes_no_unknown(application.hod_department_verified, false_label="Not Confirmed")],
        ["Finance Verification", f"{application.finance_status} | {application.finance_comments or 'N/A'}"],
        ["Not a Fee Defaulter", _yes_no_unknown(None if application.finance_fee_defaulter is None else not application.finance_fee_defaulter)],
        ["Not Availing Scholarship", _yes_no_unknown(None if application.finance_scholarship_availing is None else not application.finance_scholarship_availing)],
        ["Admin Decision", f"{application.admin_status} | {application.admin_comments or 'N/A'}"],
        ["Scholarship Status", application.scholarship_status],
        ["Rejected By Stage", application.rejected_by_stage or "N/A"],
        ["Rejection Reason", application.last_rejection_reason() or "N/A"],
        ["Resubmissions", str(application.resubmission_count)],
    ]
    elements.append(_styled_table(status_rows, widths=[150, 360]))
    elements.append(Spacer(1, 12))

    if application.field_correction_items() or application.flagged_documents():
        elements.append(Paragraph("Correction Requests", styles["Heading3"]))
        correction_rows = [["Item", "Required Update"]]
        for item in application.field_correction_items():
            correction_rows.append([item["label"], item["message"]])
        for document in application.document_correction_items():
            correction_rows.append([document["label"], document["message"]])
        elements.append(_styled_table(correction_rows, repeat_rows=1, widths=[170, 340]))
        elements.append(Spacer(1, 12))

    elements.append(Paragraph("Tracking Timeline", styles["Heading3"]))
    timeline_rows = [["Milestone", "Status", "Date", "Detail"]]
    for item in build_tracking_timeline(application):
        timeline_rows.append(
            [
                item["title"],
                item["state"].title(),
                item["timestamp"].strftime("%d %b %Y") if item["timestamp"] else "-",
                item["detail"],
            ]
        )
    elements.append(_styled_table(timeline_rows, repeat_rows=1, widths=[120, 80, 90, 220]))
    elements.append(Spacer(1, 12))

    document_rows = [["Document Type", "Status", "Original File", "Uploaded At"]]
    if application.documents:
        for uploaded_document in application.documents:
            document_rows.append(
                [
                    uploaded_document.document_type,
                    uploaded_document.review_status or "Accepted",
                    uploaded_document.original_filename,
                    uploaded_document.uploaded_at.strftime("%d %b %Y %I:%M %p"),
                ]
            )
    else:
        document_rows.append(["No documents uploaded", "-", "-", "-"])
    elements.append(_styled_table(document_rows, repeat_rows=1, widths=[120, 80, 220, 150]))
    elements.append(Spacer(1, 12))

    log_rows = [["Stage", "Action", "Actor", "Remarks", "Date"]]
    logs = sorted(application.logs, key=lambda item: item.created_at)
    for log in logs:
        log_rows.append(
            [
                log.stage,
                log.action,
                log.actor.full_name if log.actor else "-",
                log.remarks or "-",
                log.created_at.strftime("%d %b %Y %I:%M %p"),
            ]
        )
    if len(log_rows) == 1:
        log_rows.append(["No activity", "-", "-", "-", "-"])
    elements.append(_styled_table(log_rows, repeat_rows=1, widths=[110, 80, 110, 150, 110]))
    elements.append(Spacer(1, 12))

    image_documents_added = False
    for uploaded_document in application.documents:
        mime_type = (uploaded_document.mime_type or "").lower()
        resolved_document_path = resolve_managed_upload_path(uploaded_document.file_path)
        if not mime_type.startswith("image/") or not resolved_document_path:
            continue

        image_documents_added = True
        elements.append(Paragraph(f"Document Preview: {uploaded_document.document_type}", styles["Heading3"]))
        elements.append(Spacer(1, 6))
        try:
            image = ReportImage(resolved_document_path)
            max_width = 470
            max_height = 250
            width_scale = max_width / image.drawWidth
            height_scale = max_height / image.drawHeight
            scale = min(width_scale, height_scale, 1)
            image.drawWidth = image.drawWidth * scale
            image.drawHeight = image.drawHeight * scale
            elements.append(image)
        except Exception:
            elements.append(Paragraph("Preview unavailable for this image file.", styles["BodyText"]))
        elements.append(Spacer(1, 10))

    if not image_documents_added:
        elements.append(Paragraph("No image-based documents available for preview.", styles["BodyText"]))

    pdf.build(elements)
    buffer.seek(0)
    return buffer


def generate_report_pdf(payload):
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=24,
        leftMargin=24,
        topMargin=28,
        bottomMargin=24,
    )
    styles = _report_styles()
    elements = [
        Paragraph("Student Loan Management Report Pack", styles["Title"]),
        Paragraph(f"Generated on {payload['generated_at'].strftime('%d %b %Y %I:%M %p')}", styles["Meta"]),
        Spacer(1, 12),
    ]

    elements.append(Paragraph(_report_filter_label(payload["filters"]), styles["BodyText"]))
    elements.append(Spacer(1, 12))

    summary_rows = [
        ["Metric", "Value"],
        ["Total Applications", payload["summary"]["total_applications"]],
        ["Approved Loans", payload["summary"]["approved_loans"]],
        ["Rejected Loans", payload["summary"]["rejected_loans"]],
        ["Pending Applications", payload["summary"]["pending_applications"]],
    ]
    elements.append(_styled_table(summary_rows, widths=[190, 120]))
    elements.append(Spacer(1, 12))

    status_rows = [["Status", "Count"]]
    for status_name, count in payload["status_rows"]:
        status_rows.append([status_name, count])
    elements.append(_styled_table(status_rows, widths=[210, 100]))
    elements.append(Spacer(1, 12))

    department_rows = [["Department", "Applications", "Assigned HOD"]]
    coverage_map = {department.name: (department.hod.full_name if department.hod else "Unassigned") for department in payload["coverage_rows"]}
    for department_name, total in payload["department_rows"]:
        department_rows.append([department_name, total, coverage_map.get(department_name, "Unassigned")])
    if len(department_rows) == 1:
        department_rows.append(["No data", 0, "-"])
    elements.append(_styled_table(department_rows, widths=[220, 100, 180]))
    elements.append(Spacer(1, 12))

    application_rows = [["Loan ID", "Student", "Department", "Program", "Status", "Requested", "Issued"]]
    for application in payload["applications"]:
        application_rows.append(
            [
                application.loan_id,
                application.student_name,
                application.department,
                application.program_level or "BS",
                application.status,
                f"PKR {application.loan_amount_requested}",
                f"PKR {application.loan_amount_issued}" if application.loan_amount_issued else "-",
            ]
        )
    if len(application_rows) == 1:
        application_rows.append(["No applications found", "-", "-", "-", "-", "-", "-"])
    elements.append(_styled_table(application_rows, repeat_rows=1, widths=[90, 150, 150, 80, 90, 90, 90]))

    document.build(elements)
    buffer.seek(0)
    return buffer, "loan_report.pdf"


def generate_report_csv(payload):
    csv_buffer = StringIO()
    writer = csv.writer(csv_buffer)
    writer.writerow(["Student Loan Management Report"])
    writer.writerow(["Generated At", payload["generated_at"].strftime("%d %b %Y %I:%M %p")])
    writer.writerow(["Filters", _report_filter_label(payload["filters"])])
    writer.writerow([])
    writer.writerow(["Metric", "Value"])
    writer.writerow(["Total Applications", payload["summary"]["total_applications"]])
    writer.writerow(["Approved Loans", payload["summary"]["approved_loans"]])
    writer.writerow(["Rejected Loans", payload["summary"]["rejected_loans"]])
    writer.writerow(["Pending Applications", payload["summary"]["pending_applications"]])
    writer.writerow([])
    writer.writerow(["Loan ID", "Student Name", "Roll Number", "Department", "Program", "Status", "Requested", "Issued", "Student IBAN"])
    for application in payload["applications"]:
        writer.writerow(
            [
                application.loan_id,
                application.student_name,
                application.roll_number,
                application.department,
                application.program_level or "BS",
                application.status,
                application.loan_amount_requested,
                application.loan_amount_issued or "",
                application.student_iban or "",
            ]
        )

    byte_buffer = BytesIO(csv_buffer.getvalue().encode("utf-8-sig"))
    byte_buffer.seek(0)
    return byte_buffer, "loan_report.csv"


def generate_report_excel(payload):
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise RuntimeError("Excel export requires openpyxl. Install it with: pip install openpyxl") from exc

    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = "Executive Summary"
    summary_sheet.append(["Student Loan Management Report"])
    summary_sheet.append([f"Generated: {payload['generated_at'].strftime('%d %b %Y %I:%M %p')}"])
    summary_sheet.append([_report_filter_label(payload["filters"])])
    summary_sheet.append([])
    summary_sheet.append(["Metric", "Value"])
    summary_sheet.append(["Total Applications", payload["summary"]["total_applications"]])
    summary_sheet.append(["Approved Loans", payload["summary"]["approved_loans"]])
    summary_sheet.append(["Rejected Loans", payload["summary"]["rejected_loans"]])
    summary_sheet.append(["Pending Applications", payload["summary"]["pending_applications"]])

    status_sheet = workbook.create_sheet(title="Status Breakdown")
    status_sheet.append(["Status", "Count"])
    for row in payload["status_rows"]:
        status_sheet.append(list(row))

    department_sheet = workbook.create_sheet(title="Department Coverage")
    department_sheet.append(["Department", "Category", "Assigned HOD"])
    for department in payload["coverage_rows"]:
        department_sheet.append([department.name, department.category, department.hod.full_name if department.hod else "Unassigned"])

    application_sheet = workbook.create_sheet(title="Applications")
    application_sheet.append(
        [
            "Loan ID",
            "Student Name",
            "Roll Number",
            "Department",
            "Program",
            "Status",
            "Requested Amount",
            "Issued Amount",
            "Disbursement Date",
            "Repayment Deadline",
        ]
    )
    for application in payload["applications"]:
        application_sheet.append(
            [
                application.loan_id,
                application.student_name,
                application.roll_number,
                application.department,
                application.program_level or "BS",
                application.status,
                float(application.loan_amount_requested),
                float(application.loan_amount_issued) if application.loan_amount_issued else None,
                application.expected_disbursement_date.isoformat() if application.expected_disbursement_date else "",
                application.repayment_deadline.isoformat() if application.repayment_deadline else "",
            ]
        )

    header_fill = PatternFill("solid", fgColor="0F172A")
    header_font = Font(bold=True, color="FFFFFF")
    title_font = Font(size=16, bold=True)

    for worksheet in workbook.worksheets:
        for row_index, row in enumerate(worksheet.iter_rows(min_row=1, max_row=worksheet.max_row), start=1):
            if row_index == 1 and worksheet.title == "Executive Summary":
                row[0].font = title_font
            elif row_index >= 5 or worksheet.title != "Executive Summary":
                if row_index == 5 and worksheet.title == "Executive Summary" or row_index == 1 and worksheet.title != "Executive Summary":
                    for header_cell in row:
                        header_cell.font = header_font
                        header_cell.fill = header_fill
                        header_cell.alignment = Alignment(horizontal="center")
        for column_cells in worksheet.columns:
            max_length = max(len(str(cell.value or "")) for cell in column_cells)
            worksheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(max_length + 2, 12), 28)

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer, "loan_report.xlsx"


def generate_report_word(payload):
    try:
        from docx import Document as WordDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("Word export requires python-docx. Install it with: pip install python-docx") from exc

    document = WordDocument()
    title = document.add_heading("Student Loan Management Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Generated on {payload['generated_at'].strftime('%d %b %Y %I:%M %p')}")
    document.add_paragraph(_report_filter_label(payload["filters"]))

    document.add_heading("Executive Summary", level=1)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Medium Shading 1 Accent 1"
    summary_headers = summary_table.rows[0].cells
    summary_headers[0].text = "Metric"
    summary_headers[1].text = "Value"
    for metric, value in [
        ("Total Applications", payload["summary"]["total_applications"]),
        ("Approved Loans", payload["summary"]["approved_loans"]),
        ("Rejected Loans", payload["summary"]["rejected_loans"]),
        ("Pending Applications", payload["summary"]["pending_applications"]),
    ]:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = str(metric)
        row_cells[1].text = str(value)

    document.add_heading("Department Coverage", level=1)
    coverage_table = document.add_table(rows=1, cols=3)
    coverage_table.style = "Light Grid Accent 1"
    coverage_headers = coverage_table.rows[0].cells
    coverage_headers[0].text = "Department"
    coverage_headers[1].text = "Category"
    coverage_headers[2].text = "Assigned HOD"
    for department in payload["coverage_rows"]:
        cells = coverage_table.add_row().cells
        cells[0].text = department.name
        cells[1].text = department.category
        cells[2].text = department.hod.full_name if department.hod else "Unassigned"

    document.add_heading("Applications", level=1)
    applications_table = document.add_table(rows=1, cols=7)
    applications_table.style = "Light List Accent 1"
    headers = applications_table.rows[0].cells
    headers[0].text = "Loan ID"
    headers[1].text = "Student"
    headers[2].text = "Department"
    headers[3].text = "Program"
    headers[4].text = "Status"
    headers[5].text = "Requested"
    headers[6].text = "Issued"
    for application in payload["applications"]:
        cells = applications_table.add_row().cells
        cells[0].text = application.loan_id
        cells[1].text = application.student_name
        cells[2].text = application.department
        cells[3].text = application.program_level or "BS"
        cells[4].text = application.status
        cells[5].text = f"PKR {application.loan_amount_requested}"
        cells[6].text = f"PKR {application.loan_amount_issued}" if application.loan_amount_issued else "-"

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer, "loan_report.docx"


def generate_report_chart(payload):
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError as exc:
        raise RuntimeError("Chart export requires matplotlib. Install it with: pip install matplotlib") from exc

    status_labels = [name for name, count in payload["status_rows"] if count > 0]
    status_values = [count for _, count in payload["status_rows"] if count > 0]
    if not status_labels:
        status_labels = ["No Data"]
        status_values = [0]

    department_labels = [row[0] or "Unknown" for row in payload["department_rows"]]
    department_values = [row[1] for row in payload["department_rows"]]
    if not department_labels:
        department_labels = ["No Data"]
        department_values = [0]

    plt.style.use("seaborn-v0_8-whitegrid")
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    fig.patch.set_facecolor("#f8fafc")

    status_colors = ["#0f766e", "#1d4ed8", "#f59e0b", "#9333ea", "#16a34a", "#dc2626"][: len(status_labels)]
    axes[0].bar(status_labels, status_values, color=status_colors, edgecolor="#0f172a", linewidth=0.4)
    axes[0].set_title("Application Status Breakdown", fontsize=13, fontweight="bold")
    axes[0].set_ylabel("Applications")
    axes[0].tick_params(axis="x", rotation=25)
    for index, value in enumerate(status_values):
        axes[0].text(index, value + 0.05, str(value), ha="center", va="bottom", fontsize=9)

    axes[1].barh(department_labels, department_values, color="#2563eb")
    axes[1].set_title("Department-wise Applications", fontsize=13, fontweight="bold")
    axes[1].set_xlabel("Applications")
    for index, value in enumerate(department_values):
        axes[1].text(value + 0.05, index, str(value), va="center", fontsize=9)

    fig.suptitle("Student Loan Management Analytics", fontsize=16, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.95])

    buffer = BytesIO()
    fig.savefig(buffer, format="png", dpi=180, bbox_inches="tight")
    plt.close(fig)
    buffer.seek(0)
    return buffer, "loan_report_chart.png"


def _report_styles():
    styles = getSampleStyleSheet()
    styles["Title"].textColor = colors.HexColor("#0f172a")
    styles["Heading2"].textColor = colors.HexColor("#0f1d33")
    styles["Heading3"].textColor = colors.HexColor("#0f766e")
    styles.add(
        ParagraphStyle(
            name="Meta",
            parent=styles["BodyText"],
            fontSize=9,
            textColor=colors.HexColor("#475569"),
            leading=12,
        )
    )
    return styles


def _yes_no_unknown(value, true_label="Yes", false_label="No", unknown_label="Not checked"):
    if value is True:
        return true_label
    if value is False:
        return false_label
    return unknown_label


def _hod_validation_summary(application):
    return (
        f"CGPA check: {_yes_no_unknown(application.hod_cgpa_verified, false_label='Not verified')} | "
        f"Department check: {_yes_no_unknown(application.hod_department_verified, false_label='Not verified')}"
    )


def _finance_validation_summary(application):
    return (
        f"Not a fee defaulter: {_yes_no_unknown(None if application.finance_fee_defaulter is None else not application.finance_fee_defaulter)} | "
        f"Not availing scholarship: {_yes_no_unknown(None if application.finance_scholarship_availing is None else not application.finance_scholarship_availing)}"
    )


def _normalized_timeline_timestamp(value):
    if value is None:
        return None
    if isinstance(value, datetime):
        return value
    if isinstance(value, date):
        return datetime.combine(value, datetime.min.time())
    return None


def _format_currency(value):
    if value is None:
        return "PKR 0.00"
    return f"PKR {float(value):,.2f}"


def _styled_table(rows, repeat_rows=0, widths=None):
    table = Table(rows, repeatRows=repeat_rows, colWidths=widths)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEADING", (0, 0), (-1, -1), 13),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def _report_filter_label(filters):
    active_filters = []
    if filters.get("status"):
        active_filters.append(f"Status: {filters['status']}")
    if filters.get("department"):
        active_filters.append(f"Department: {filters['department']}")
    return "Filters applied - " + ", ".join(active_filters) if active_filters else "Filters applied - All applications"
